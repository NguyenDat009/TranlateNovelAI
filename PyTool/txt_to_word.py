#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tool chuyển đổi file .txt sang .docx (Word)
Tự động sử dụng tên file làm tiêu đề
Hỗ trợ xử lý nhiều file cùng lúc
"""

import os
import glob
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def detect_encoding(file_path):
    """
    Phát hiện encoding của file
    """
    # Kiểm tra BOM trước
    with open(file_path, 'rb') as f:
        raw = f.read(4)
        if raw.startswith(b'\xff\xfe\x00\x00'):
            return 'utf-32-le'
        elif raw.startswith(b'\x00\x00\xfe\xff'):
            return 'utf-32-be'
        elif raw.startswith(b'\xff\xfe') or raw.startswith(b'\xfe\xff'):
            return 'utf-16'
        elif raw.startswith(b'\xef\xbb\xbf'):
            return 'utf-8-sig'
    
    # Thử các encoding phổ biến
    encodings = ['utf-8', 'utf-8-sig', 'utf-16', 'utf-16-le', 'utf-16-be', 
                 'cp1252', 'latin-1', 'ascii']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                sample = f.read(8192)
                if sample:
                    return encoding
        except (UnicodeDecodeError, UnicodeError, LookupError):
            continue
    
    return 'utf-8'


def txt_to_word(txt_path, output_folder=None):
    """
    Chuyển đổi file .txt sang .docx
    
    Args:
        txt_path (str): Đường dẫn đến file .txt
        output_folder (str): Thư mục lưu file output (None = cùng thư mục với file gốc)
    
    Returns:
        tuple: (success, docx_path, error_message)
    """
    try:
        # Kiểm tra file tồn tại
        if not os.path.exists(txt_path):
            return False, None, f"File không tồn tại: {txt_path}"
        
        # Lấy tên file (không có đuôi)
        file_name = os.path.splitext(os.path.basename(txt_path))[0]
        
        # Xác định đường dẫn output
        if output_folder:
            os.makedirs(output_folder, exist_ok=True)
            docx_path = os.path.join(output_folder, f"{file_name}.docx")
        else:
            docx_path = os.path.join(os.path.dirname(txt_path), f"{file_name}.docx")
        
        # Đọc nội dung file
        encoding = detect_encoding(txt_path)
        with open(txt_path, 'r', encoding=encoding, errors='replace') as f:
            content = f.read()
        
        # Xử lý nội dung
        if content.startswith('\ufeff'):
            content = content[1:]
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        content = content.strip()
        
        if not content:
            return False, None, "File rỗng"
        
        # Tạo document Word
        doc = Document()
        
        # Cài đặt lề
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Thêm tiêu đề (tên file)
        title = doc.add_heading(file_name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm một dòng trống
        doc.add_paragraph()
        
        # Tách nội dung thành các đoạn văn
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # Kiểm tra xem có phải là tiêu đề không (dòng ngắn, không có dấu câu)
                lines = para_text.strip().split('\n')
                
                # Nếu là dòng đơn và ngắn (< 100 ký tự), có thể là tiêu đề phụ
                if len(lines) == 1 and len(lines[0]) < 100 and not lines[0].endswith(('.', ',', '!', '?', ';')):
                    # Thêm như heading level 2
                    doc.add_heading(lines[0], level=2)
                else:
                    # Thêm như đoạn văn bình thường
                    para = doc.add_paragraph()
                    
                    # Giữ nguyên các dòng ngắt trong đoạn văn
                    for i, line in enumerate(lines):
                        if line.strip():
                            if i > 0:
                                para.add_run('\n')
                            run = para.add_run(line.strip())
                            run.font.size = Pt(12)
                    
                    # Định dạng đoạn văn
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.space_after = Pt(6)
                    para.paragraph_format.first_line_indent = Inches(0.5)
        
        # Lưu file
        doc.save(docx_path)
        
        return True, docx_path, None
        
    except Exception as e:
        return False, None, str(e)


def convert_multiple_files(input_path, output_folder=None):
    """
    Convert nhiều file txt sang word
    
    Args:
        input_path (str): Đường dẫn đến file hoặc thư mục
        output_folder (str): Thư mục lưu output
    """
    print("="*80)
    print(" "*20 + "CHUYỂN ĐỔI TXT SANG WORD")
    print("="*80)
    print()
    
    # Xác định danh sách file cần convert
    txt_files = []
    
    if os.path.isfile(input_path):
        # Nếu là file đơn
        if input_path.endswith('.txt'):
            txt_files = [input_path]
        else:
            print("❌ File phải có đuôi .txt!")
            return
    elif os.path.isdir(input_path):
        # Nếu là thư mục, lấy tất cả file .txt
        txt_files = sorted(glob.glob(os.path.join(input_path, '*.txt')))
    else:
        print(f"❌ Đường dẫn không hợp lệ: {input_path}")
        return
    
    if not txt_files:
        print("❌ Không tìm thấy file .txt nào!")
        return
    
    print(f"📁 Tìm thấy {len(txt_files)} file .txt")
    print(f"📝 Đang chuyển đổi...\n")
    
    # Convert từng file
    success_count = 0
    error_count = 0
    results = []
    
    for i, txt_file in enumerate(txt_files, 1):
        file_name = os.path.basename(txt_file)
        print(f"[{i}/{len(txt_files)}] Đang xử lý: {file_name}", end=' ... ')
        
        success, docx_path, error = txt_to_word(txt_file, output_folder)
        
        if success:
            print(f"✅ Thành công")
            success_count += 1
            results.append((file_name, docx_path, True))
        else:
            print(f"❌ Lỗi: {error}")
            error_count += 1
            results.append((file_name, None, False))
    
    # Hiển thị kết quả
    print(f"\n{'='*80}")
    print(f"✨ Hoàn thành!")
    print(f"📊 Thống kê:")
    print(f"   - Thành công: {success_count} file")
    print(f"   - Lỗi: {error_count} file")
    
    if output_folder:
        print(f"📁 Thư mục output: {os.path.abspath(output_folder)}")
    
    print(f"\n📄 Danh sách file đã tạo:")
    for file_name, docx_path, success in results:
        if success:
            print(f"   ✅ {file_name} → {os.path.basename(docx_path)}")
    
    if error_count > 0:
        print(f"\n❌ File lỗi:")
        for file_name, _, success in results:
            if not success:
                print(f"   ❌ {file_name}")
    
    print(f"{'='*80}")


def main():
    """
    Hàm chính
    """
    print("\n" + "="*80)
    print(" "*25 + "TXT TO WORD CONVERTER")
    print("="*80)
    print()
    print("Công cụ chuyển đổi file TXT sang DOCX (Word)")
    print("Tự động sử dụng tên file làm tiêu đề")
    print()
    
    # Nhập đường dẫn
    input_path = input("📁 Nhập đường dẫn file .txt hoặc thư mục chứa file .txt: ").strip()
    input_path = input_path.strip('"').strip("'")
    
    if not os.path.exists(input_path):
        print(f"\n❌ Đường dẫn không tồn tại: {input_path}")
        input("\nNhấn Enter để thoát...")
        return
    
    # Hỏi thư mục output
    output_folder = input("📂 Nhập thư mục lưu file Word (Enter để lưu cùng thư mục với file gốc): ").strip()
    output_folder = output_folder.strip('"').strip("'")
    
    if not output_folder:
        output_folder = None
        print("💡 Sẽ lưu file Word cùng thư mục với file TXT gốc")
    
    print()
    
    # Chuyển đổi
    convert_multiple_files(input_path, output_folder)


if __name__ == '__main__':
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⚠️ Đã dừng chương trình!")
    except Exception as e:
        print(f"\n❌ Lỗi: {str(e)}")
        import traceback
        traceback.print_exc()
    
    # Giữ cửa sổ console mở
    input("\nNhấn Enter để thoát...")

