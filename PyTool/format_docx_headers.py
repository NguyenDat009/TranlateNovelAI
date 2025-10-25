"""
Script để format lại headers trong file DOCX:
- Giữ nguyên Title ở Header 1
- Chuyển các "Chương" từ Header 1 sang Header 2
- Thêm page break trước mỗi chương nếu chưa có
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
import re
import os
from pathlib import Path


def is_chapter_heading(text):
    """
    Kiểm tra xem đoạn text có phải là tiêu đề chương không.
    Các pattern phổ biến:
    - Chương 1
    - Chương I
    - Chương 001
    - Chapter 1
    - CHƯƠNG 1
    """
    if not text:
        return False
    
    text = text.strip()
    
    # Pattern cho các dạng chương
    patterns = [
        r'^Chương\s+\d+',  # Chương 1, Chương 2, ...
        r'^CHƯƠNG\s+\d+',  # CHƯƠNG 1, CHƯƠNG 2, ...
        r'^Chương\s+[IVXLCDM]+',  # Chương I, Chương II, ... (số La Mã)
        r'^CHƯƠNG\s+[IVXLCDM]+',  # CHƯƠNG I, CHƯƠNG II, ...
        r'^Chapter\s+\d+',  # Chapter 1, Chapter 2, ...
        r'^CHAPTER\s+\d+',  # CHAPTER 1, CHAPTER 2, ...
    ]
    
    for pattern in patterns:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    
    return False


def has_page_break_before(paragraph):
    """
    Kiểm tra xem paragraph có page break trước nó không.
    """
    # Kiểm tra trong runs của paragraph
    for run in paragraph.runs:
        if run._element.xml.find('w:br') != -1:
            if 'w:type="page"' in run._element.xml:
                return True
    
    # Kiểm tra page break trong paragraph properties
    pPr = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if pPr is not None:
        page_break = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pageBreakBefore')
        if page_break is not None:
            return True
    
    return False


def add_page_break_before(paragraph):
    """
    Thêm page break trước paragraph.
    """
    # Tạo run mới ở đầu paragraph với page break
    run = paragraph.insert_paragraph_before().add_run()
    run.add_break(WD_BREAK.PAGE)


def format_docx_headers(input_file, output_file=None):
    """
    Format lại headers trong file DOCX:
    - Giữ nguyên Title ở Header 1
    - Chuyển các "Chương" từ Header 1 sang Header 2
    - Thêm page break trước mỗi chương
    
    Args:
        input_file: Đường dẫn file DOCX đầu vào
        output_file: Đường dẫn file DOCX đầu ra (nếu None, sẽ ghi đè file gốc)
    """
    print(f"Đang xử lý file: {input_file}")
    
    # Mở document
    doc = Document(input_file)
    
    # Thống kê
    stats = {
        'total_paragraphs': len(doc.paragraphs),
        'header1_found': 0,
        'chapters_converted': 0,
        'page_breaks_added': 0,
        'titles_kept': 0
    }
    
    # Duyệt qua tất cả các paragraph
    previous_para = None
    for i, para in enumerate(doc.paragraphs):
        # Kiểm tra nếu là Header 1
        if para.style.name == 'Heading 1':
            stats['header1_found'] += 1
            text = para.text.strip()
            
            # Kiểm tra nếu là chương
            if is_chapter_heading(text):
                print(f"  Tìm thấy chương: {text}")
                
                # Chuyển sang Header 2
                para.style = 'Heading 2'
                stats['chapters_converted'] += 1
                
                # Kiểm tra và thêm page break nếu chưa có
                # (không thêm page break cho chương đầu tiên)
                if i > 0 and not has_page_break_before(para):
                    # Kiểm tra paragraph trước có phải là page break không
                    if previous_para and previous_para.text.strip() == '':
                        # Kiểm tra xem có runs với page break không
                        has_break = False
                        for run in previous_para.runs:
                            if 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
                                has_break = True
                                break
                        
                        if not has_break:
                            add_page_break_before(para)
                            stats['page_breaks_added'] += 1
                            print(f"    → Đã thêm page break")
                    else:
                        add_page_break_before(para)
                        stats['page_breaks_added'] += 1
                        print(f"    → Đã thêm page break")
                
                print(f"    → Đã chuyển sang Header 2")
            else:
                # Không phải chương, giữ nguyên Header 1 (Title)
                stats['titles_kept'] += 1
                print(f"  Giữ nguyên title: {text}")
        
        previous_para = para
    
    # Xác định file output
    if output_file is None:
        output_file = input_file
    
    # Lưu document
    doc.save(output_file)
    
    print(f"\n✓ Hoàn thành!")
    print(f"  File output: {output_file}")
    print(f"\nThống kê:")
    print(f"  - Tổng số paragraphs: {stats['total_paragraphs']}")
    print(f"  - Header 1 tìm thấy: {stats['header1_found']}")
    print(f"  - Chương đã chuyển sang Header 2: {stats['chapters_converted']}")
    print(f"  - Page breaks đã thêm: {stats['page_breaks_added']}")
    print(f"  - Titles giữ nguyên Header 1: {stats['titles_kept']}")


def main():
    """
    Hàm main để chạy script.
    """
    print("=" * 70)
    print("SCRIPT FORMAT LẠI HEADERS TRONG FILE DOCX")
    print("=" * 70)
    print()
    
    # Nhập đường dẫn file
    input_file = input("Nhập đường dẫn file DOCX cần format: ").strip().strip('"')
    
    # Kiểm tra file tồn tại
    if not os.path.exists(input_file):
        print(f"❌ Lỗi: Không tìm thấy file '{input_file}'")
        return
    
    # Kiểm tra extension
    if not input_file.lower().endswith('.docx'):
        print(f"❌ Lỗi: File phải có định dạng .docx")
        return
    
    # Hỏi xem có muốn tạo file mới không
    create_new = input("\nBạn có muốn tạo file mới? (y/n, mặc định là ghi đè file gốc): ").strip().lower()
    
    output_file = None
    if create_new == 'y':
        # Tạo tên file mới
        path = Path(input_file)
        output_file = str(path.parent / f"{path.stem}_formatted{path.suffix}")
        print(f"File mới sẽ được lưu tại: {output_file}")
    else:
        print("File gốc sẽ bị ghi đè.")
        confirm = input("Bạn có chắc chắn? (y/n): ").strip().lower()
        if confirm != 'y':
            print("Đã hủy.")
            return
    
    print()
    
    # Xử lý file
    try:
        format_docx_headers(input_file, output_file)
    except Exception as e:
        print(f"\n❌ Lỗi khi xử lý file: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
